import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

def add_code_to_doc(doc, title, code):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_page_break()

def main():
    base_dir = r"C:\Users\josem\Downloads\cpsl-campus-interactivo\src"
    output_path = r"C:\Users\josem\Downloads\cpsl-campus-interactivo\Auditoria_Codigo_Completo.docx"
    
    doc = Document()
    doc.add_heading('Auditoría Código Completo - Campus Interactivo', 0)
    
    allowed_extensions = {'.js', '.jsx', '.css'}
    
    for root, dirs, files in os.walk(base_dir):
        for file in files:
            ext = os.path.splitext(file)[1]
            if ext in allowed_extensions:
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, base_dir)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code = f.read()
                    
                    add_code_to_doc(doc, f"Archivo: src/{rel_path.replace(os.sep, '/')}", code)
                except Exception as e:
                    print(f"No se pudo leer {file_path}: {e}")
                    
    # Agregamos App.jsx, index.css, main.jsx, index.html si no están
    root_files = ['App.jsx', 'index.css', 'main.jsx']
    for rf in root_files:
        path = os.path.join(base_dir, rf)
        if os.path.exists(path):
            pass # Ya fue incluido por os.walk
            
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    main()
